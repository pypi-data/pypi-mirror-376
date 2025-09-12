"""Microsoft Office document loaders - PowerPoint, Word, Excel."""

import shutil

from ... import matchers
from ...core import Attachment, loader


@loader(match=matchers.pptx_match)
def pptx_to_python_pptx(att: Attachment) -> Attachment:
    """Load PowerPoint using python-pptx with automatic input source handling."""
    try:
        from pptx import Presentation

        # Use the new input_source property - no more repetitive patterns!
        att._obj = Presentation(att.input_source)

    except ImportError as err:
        raise ImportError(
            "python-pptx is required for PowerPoint loading. Install with: pip install python-pptx"
        ) from err
    return att


@loader(match=matchers.docx_match)
def docx_to_python_docx(att: Attachment) -> Attachment:
    """Load Word document using python-docx with automatic input source handling."""
    try:
        from docx import Document

        # Use the new input_source property - no more repetitive patterns!
        att._obj = Document(att.input_source)

    except ImportError as err:
        raise ImportError(
            "python-docx is required for Word document loading. Install with: pip install python-docx"
        ) from err
    return att


@loader(match=matchers.excel_match)
def excel_to_openpyxl(att: Attachment) -> Attachment:
    """Load Excel workbook using openpyxl with automatic input source handling."""
    try:
        from openpyxl import load_workbook

        # Use the new input_source property - no more repetitive patterns!
        att._obj = load_workbook(att.input_source, read_only=True)

    except ImportError as err:
        raise ImportError(
            "openpyxl is required for Excel loading. Install with: pip install openpyxl"
        ) from err
    return att


class LibreOfficeDocument:
    """A proxy object representing a document to be handled by LibreOffice."""

    def __init__(self, path: str):
        self.path = path

    def __repr__(self):
        return f"LibreOfficeDocument(path='{self.path}')"


@loader(match=matchers.excel_match)
def excel_to_libreoffice(att: Attachment) -> Attachment:
    """
    Prepares an Excel file for processing via LibreOffice by checking for the
    binary and wrapping the attachment in a LibreOfficeDocument proxy object.
    """
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice/soffice not found. This loader requires a LibreOffice installation."
        )

    # Store the binary path for the presenter to use
    att.metadata["libreoffice_binary_path"] = soffice

    # Set the object to our proxy type for dispatch
    att._obj = LibreOfficeDocument(att.path)

    return att
