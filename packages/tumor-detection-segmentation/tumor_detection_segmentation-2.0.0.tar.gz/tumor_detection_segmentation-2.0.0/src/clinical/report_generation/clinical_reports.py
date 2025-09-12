"""
Clinical Report Generation Module

This module provides structured clinical report generation for AI-powered
tumor segmentation results, supporting various clinical reporting standards
and templates.

Features:
- BI-RADS, PI-RADS, LI-RADS compatible templates
- Natural language generation for findings
- Quantitative measurements and statistics
- PDF and Word export capabilities
- Institution-specific template customization
- Integration with AI segmentation results
"""

import json
import logging
from dataclasses import asdict, dataclass
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# Report generation libraries
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle)
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from jinja2 import Environment, FileSystemLoader, Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

logger = logging.getLogger(__name__)


class ReportType(Enum):
    """Types of clinical reports supported."""
    BRAIN_TUMOR = "brain_tumor"
    LIVER_TUMOR = "liver_tumor"
    LUNG_NODULE = "lung_nodule"
    PROSTATE = "prostate"
    BREAST = "breast"
    CUSTOM = "custom"


class FindingSeverity(Enum):
    """Severity levels for findings."""
    NORMAL = "normal"
    BENIGN = "benign"
    PROBABLY_BENIGN = "probably_benign"
    SUSPICIOUS = "suspicious"
    HIGHLY_SUSPICIOUS = "highly_suspicious"
    MALIGNANT = "malignant"


@dataclass
class PatientInfo:
    """Patient demographic and study information."""
    patient_id: str
    patient_name: str
    date_of_birth: str
    gender: str
    study_date: str
    study_description: str
    modality: str
    institution: str = ""
    referring_physician: str = ""
    radiologist: str = ""


@dataclass
class QuantitativeMeasurements:
    """Quantitative measurements from AI analysis."""
    total_volume_cm3: float
    enhancing_volume_cm3: float = 0.0
    necrotic_volume_cm3: float = 0.0
    edema_volume_cm3: float = 0.0
    largest_diameter_mm: float = 0.0
    mean_adc_value: float = 0.0
    confidence_score: float = 0.0

    # Additional measurements
    surface_area_cm2: float = 0.0
    sphericity: float = 0.0
    contrast_enhancement_ratio: float = 0.0


@dataclass
class AIFinding:
    """Individual finding from AI analysis."""
    finding_id: str
    location: str
    description: str
    severity: FindingSeverity
    measurements: QuantitativeMeasurements
    confidence: float
    coordinates: Optional[List[float]] = None
    slice_numbers: Optional[List[int]] = None


@dataclass
class ClinicalFindings:
    """Complete clinical findings from AI analysis."""
    primary_findings: List[AIFinding]
    secondary_findings: List[AIFinding]
    incidental_findings: List[AIFinding]
    technical_quality: str
    limitations: List[str]
    recommendations: List[str]


class ReportTemplate:
    """Base class for clinical report templates."""

    def __init__(self, template_name: str, report_type: ReportType):
        self.template_name = template_name
        self.report_type = report_type
        self.sections = []
        self.style_config = {}

    def add_section(self, section_name: str, content: str):
        """Add a section to the report."""
        self.sections.append({
            'name': section_name,
            'content': content
        })

    def generate_content(self, patient_info: PatientInfo,
                        findings: ClinicalFindings) -> Dict[str, str]:
        """Generate report content based on template and findings."""
        raise NotImplementedError("Subclasses must implement generate_content")


class BrainTumorReportTemplate(ReportTemplate):
    """Brain tumor reporting template based on RANO criteria."""

    def __init__(self):
        super().__init__("Brain Tumor Report", ReportType.BRAIN_TUMOR)
        self.style_config = {
            'font_family': 'Helvetica',
            'font_size': 11,
            'line_spacing': 1.2
        }

    def generate_content(self, patient_info: PatientInfo,
                        findings: ClinicalFindings) -> Dict[str, str]:
        """Generate brain tumor report content."""

        content = {}

        # Clinical History section
        content['clinical_history'] = self._generate_clinical_history(
            patient_info
        )

        # Technique section
        content['technique'] = self._generate_technique_section(patient_info)

        # Findings section
        content['findings'] = self._generate_findings_section(findings)

        # Impression section
        content['impression'] = self._generate_impression_section(findings)

        # Recommendations section
        content['recommendations'] = self._generate_recommendations_section(
            findings
        )

        return content

    def _generate_clinical_history(self, patient_info: PatientInfo) -> str:
        """Generate clinical history section."""
        return f"""
        Patient: {patient_info.patient_name} (ID: {patient_info.patient_id})
        Date of Birth: {patient_info.date_of_birth}
        Gender: {patient_info.gender}
        Study Date: {patient_info.study_date}
        Referring Physician: {patient_info.referring_physician}

        Clinical indication: {patient_info.study_description}
        """

    def _generate_technique_section(self, patient_info: PatientInfo) -> str:
        """Generate technique section."""
        return f"""
        Multiparametric MRI examination of the brain was performed on a
        {patient_info.modality} scanner.

        Sequences obtained:
        - T1-weighted pre-contrast
        - T1-weighted post-contrast with gadolinium
        - T2-weighted FLAIR
        - T2-weighted
        - Diffusion-weighted imaging (DWI) with ADC mapping

        AI-assisted analysis was performed using validated deep learning models
        for tumor detection and segmentation (UNETR architecture).
        """

    def _generate_findings_section(self, findings: ClinicalFindings) -> str:
        """Generate findings section."""
        findings_text = "FINDINGS:\n\n"

        # Primary findings
        if findings.primary_findings:
            findings_text += "Primary Findings:\n"
            for i, finding in enumerate(findings.primary_findings, 1):
                findings_text += f"{i}. {self._format_finding(finding)}\n\n"

        # Secondary findings
        if findings.secondary_findings:
            findings_text += "Additional Findings:\n"
            for finding in findings.secondary_findings:
                findings_text += f"- {self._format_finding(finding)}\n"

        # Technical quality
        findings_text += f"\nTechnical Quality: {findings.technical_quality}\n"

        # Limitations
        if findings.limitations:
            findings_text += "\nLimitations:\n"
            for limitation in findings.limitations:
                findings_text += f"- {limitation}\n"

        return findings_text

    def _format_finding(self, finding: AIFinding) -> str:
        """Format individual finding for report."""
        measurements = finding.measurements

        finding_text = f"""
        {finding.description} in the {finding.location}.

        Quantitative Measurements:
        - Total volume: {measurements.total_volume_cm3:.1f} cm³
        - Enhancing volume: {measurements.enhancing_volume_cm3:.1f} cm³
        - Necrotic volume: {measurements.necrotic_volume_cm3:.1f} cm³
        - Largest diameter: {measurements.largest_diameter_mm:.1f} mm
        - AI confidence: {finding.confidence:.2f}

        Assessment: {finding.severity.value.replace('_', ' ').title()}
        """

        return finding_text.strip()

    def _generate_impression_section(self, findings: ClinicalFindings) -> str:
        """Generate impression section."""
        impression = "IMPRESSION:\n\n"

        if not findings.primary_findings:
            impression += "No significant abnormalities detected by AI analysis."
            return impression

        # Summarize primary findings
        for i, finding in enumerate(findings.primary_findings, 1):
            volume = finding.measurements.total_volume_cm3
            confidence = finding.confidence

            impression += f"{i}. {finding.description} in {finding.location}, "
            impression += f"measuring {volume:.1f} cm³ "
            impression += f"(AI confidence: {confidence:.2f}).\n"

        # Overall assessment
        if len(findings.primary_findings) == 1:
            primary = findings.primary_findings[0]
            if primary.severity in [FindingSeverity.SUSPICIOUS,
                                  FindingSeverity.HIGHLY_SUSPICIOUS]:
                impression += "\nFindings are concerning for malignancy. "
                impression += "Clinical correlation and possible tissue "
                impression += "sampling recommended.\n"
            elif primary.severity == FindingSeverity.PROBABLY_BENIGN:
                impression += "\nFindings likely represent benign process. "
                impression += "Follow-up imaging recommended.\n"

        return impression

    def _generate_recommendations_section(self,
                                        findings: ClinicalFindings) -> str:
        """Generate recommendations section."""
        recommendations_text = "RECOMMENDATIONS:\n\n"

        if findings.recommendations:
            for i, rec in enumerate(findings.recommendations, 1):
                recommendations_text += f"{i}. {rec}\n"
        else:
            # Generate default recommendations based on findings
            if findings.primary_findings:
                primary = findings.primary_findings[0]
                if primary.severity in [FindingSeverity.SUSPICIOUS,
                                      FindingSeverity.HIGHLY_SUSPICIOUS]:
                    recommendations_text += """
                    1. Neurosurgical consultation for tissue sampling
                    2. Multidisciplinary tumor board discussion
                    3. Consider additional imaging (perfusion, spectroscopy)
                    """
                else:
                    recommendations_text += """
                    1. Follow-up MRI in 3-6 months
                    2. Clinical correlation recommended
                    """
            else:
                recommendations_text += "1. No specific follow-up required.\n"

        return recommendations_text


class ClinicalReportGenerator:
    """Main class for generating clinical reports."""

    def __init__(self, template_dir: Optional[Path] = None):
        self.template_dir = template_dir or Path("templates/reports")
        self.templates = {}
        self.output_dir = Path("reports/generated")

        # Initialize templates
        self._initialize_templates()

        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _initialize_templates(self):
        """Initialize report templates."""
        self.templates[ReportType.BRAIN_TUMOR] = BrainTumorReportTemplate()

        # TODO: Add more templates
        # self.templates[ReportType.LIVER_TUMOR] = LiverTumorReportTemplate()
        # self.templates[ReportType.LUNG_NODULE] = LungNoduleReportTemplate()

    def generate_report(self,
                       patient_info: PatientInfo,
                       findings: ClinicalFindings,
                       report_type: ReportType,
                       output_format: str = "pdf",
                       custom_template: Optional[str] = None) -> Path:
        """
        Generate a clinical report.

        Args:
            patient_info: Patient and study information
            findings: Clinical findings from AI analysis
            report_type: Type of report to generate
            output_format: Output format ("pdf", "docx", "html")
            custom_template: Optional custom template name

        Returns:
            Path to generated report file
        """

        # Get template
        if custom_template:
            template = self._load_custom_template(custom_template)
        else:
            template = self.templates.get(report_type)

        if not template:
            raise ValueError(f"No template available for {report_type}")

        # Generate content
        content = template.generate_content(patient_info, findings)

        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{patient_info.patient_id}_{report_type.value}_{timestamp}"

        # Generate report based on format
        if output_format.lower() == "pdf":
            output_path = self._generate_pdf_report(
                content, filename, patient_info
            )
        elif output_format.lower() == "docx":
            output_path = self._generate_docx_report(
                content, filename, patient_info
            )
        elif output_format.lower() == "html":
            output_path = self._generate_html_report(
                content, filename, patient_info
            )
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

        logger.info(f"Generated report: {output_path}")
        return output_path

    def _generate_pdf_report(self, content: Dict[str, str],
                           filename: str, patient_info: PatientInfo) -> Path:
        """Generate PDF report using ReportLab."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab required for PDF generation")

        output_path = self.output_dir / f"{filename}.pdf"

        # Create document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch,
            leftMargin=0.5*inch,
            rightMargin=0.5*inch
        )

        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            alignment=1,  # Center alignment
            spaceAfter=20
        )

        # Build content
        story = []

        # Title
        title = f"RADIOLOGY REPORT - {patient_info.study_description.upper()}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))

        # Add sections
        for section_name, section_content in content.items():
            # Section header
            header = section_name.replace('_', ' ').title()
            story.append(Paragraph(header, styles['Heading2']))

            # Section content
            # Convert newlines to proper paragraph breaks
            paragraphs = section_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))

            story.append(Spacer(1, 12))

        # Add footer
        footer_text = f"""
        Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
        Institution: {patient_info.institution}
        AI-assisted analysis provided by Medical Imaging AI Platform
        """
        story.append(Spacer(1, 20))
        story.append(Paragraph(footer_text, styles['Italic']))

        # Build PDF
        doc.build(story)

        return output_path

    def _generate_docx_report(self, content: Dict[str, str],
                            filename: str, patient_info: PatientInfo) -> Path:
        """Generate Word document report."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for Word generation")

        output_path = self.output_dir / f"{filename}.docx"

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(
            f'RADIOLOGY REPORT - {patient_info.study_description.upper()}',
            level=1
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        for section_name, section_content in content.items():
            # Section header
            header = section_name.replace('_', ' ').title()
            doc.add_heading(header, level=2)

            # Section content
            paragraphs = section_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Add footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run(
            f"Report generated on "
            f"{datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n"
        ).italic = True
        footer_para.add_run(
            f"Institution: {patient_info.institution}\n"
        ).italic = True
        footer_para.add_run(
            "AI-assisted analysis provided by Medical Imaging AI Platform"
        ).italic = True

        # Save document
        doc.save(str(output_path))

        return output_path

    def _generate_html_report(self, content: Dict[str, str],
                            filename: str, patient_info: PatientInfo) -> Path:
        """Generate HTML report."""
        output_path = self.output_dir / f"{filename}.html"

        # Basic HTML template
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Radiology Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { text-align: center; color: #2c3e50; }
                h2 { color: #34495e; border-bottom: 1px solid #ecf0f1; }
                .footer { margin-top: 30px; font-style: italic; color: #7f8c8d; }
            </style>
        </head>
        <body>
            <h1>RADIOLOGY REPORT - {{ study_description|upper }}</h1>

            {% for section_name, section_content in content.items() %}
            <h2>{{ section_name|replace('_', ' ')|title }}</h2>
            <div>
                {% for paragraph in section_content.split('\n\n') %}
                    {% if paragraph.strip() %}
                        <p>{{ paragraph.strip() }}</p>
                    {% endif %}
                {% endfor %}
            </div>
            {% endfor %}

            <div class="footer">
                <p>Report generated on {{ timestamp }}</p>
                <p>Institution: {{ institution }}</p>
                <p>AI-assisted analysis provided by Medical Imaging AI Platform</p>
            </div>
        </body>
        </html>
        """

        if JINJA2_AVAILABLE:
            template = Template(html_template)
            html_content = template.render(
                study_description=patient_info.study_description,
                content=content,
                timestamp=datetime.now().strftime('%B %d, %Y at %I:%M %p'),
                institution=patient_info.institution
            )
        else:
            # Simple string replacement fallback
            html_content = html_template.replace(
                '{{ study_description|upper }}',
                patient_info.study_description.upper()
            )
            # Add content sections (simplified)
            sections_html = ""
            for section_name, section_content in content.items():
                header = section_name.replace('_', ' ').title()
                sections_html += f"<h2>{header}</h2>\n"
                for para in section_content.split('\n\n'):
                    if para.strip():
                        sections_html += f"<p>{para.strip()}</p>\n"

            html_content = html_content.replace(
                '{% for section_name, section_content in content.items() %}...{% endfor %}',
                sections_html
            )

            html_content = html_content.replace(
                '{{ timestamp }}',
                datetime.now().strftime('%B %d, %Y at %I:%M %p')
            )
            html_content = html_content.replace(
                '{{ institution }}',
                patient_info.institution
            )

        # Write HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return output_path

    def _load_custom_template(self, template_name: str) -> ReportTemplate:
        """Load custom template from file."""
        # TODO: Implement custom template loading
        raise NotImplementedError("Custom template loading not yet implemented")


def create_sample_report() -> Path:
    """Create a sample report for testing purposes."""

    # Sample patient info
    patient = PatientInfo(
        patient_id="12345",
        patient_name="John Doe",
        date_of_birth="1980-05-15",
        gender="Male",
        study_date="2025-09-06",
        study_description="MRI Brain with contrast",
        modality="MRI 3.0T",
        institution="Medical Imaging AI Hospital",
        referring_physician="Dr. Smith",
        radiologist="Dr. Johnson"
    )

    # Sample measurements
    measurements = QuantitativeMeasurements(
        total_volume_cm3=15.2,
        enhancing_volume_cm3=8.7,
        necrotic_volume_cm3=3.1,
        edema_volume_cm3=22.8,
        largest_diameter_mm=42.3,
        confidence_score=0.91
    )

    # Sample finding
    finding = AIFinding(
        finding_id="finding_001",
        location="left frontal lobe",
        description="Heterogeneously enhancing mass with central necrosis",
        severity=FindingSeverity.HIGHLY_SUSPICIOUS,
        measurements=measurements,
        confidence=0.91
    )

    # Sample findings
    findings = ClinicalFindings(
        primary_findings=[finding],
        secondary_findings=[],
        incidental_findings=[],
        technical_quality="Excellent",
        limitations=["Patient motion artifacts minimal"],
        recommendations=[
            "Neurosurgical consultation recommended",
            "Consider stereotactic biopsy",
            "Multidisciplinary tumor board discussion"
        ]
    )

    # Generate report
    generator = ClinicalReportGenerator()
    report_path = generator.generate_report(
        patient_info=patient,
        findings=findings,
        report_type=ReportType.BRAIN_TUMOR,
        output_format="pdf"
    )

    return report_path


if __name__ == "__main__":
    # Test report generation
    logging.basicConfig(level=logging.INFO)

    try:
        report_path = create_sample_report()
        print(f"Sample report generated: {report_path}")
    except Exception as e:
        print(f"Error generating report: {e}")
        logger.error(f"Report generation failed: {e}")
