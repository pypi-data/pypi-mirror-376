from pathlib import Path
from typing import Any, Self

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from yaml import MappingNode

from tesorotools.utils.template import TemplateLoader

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _style_container_table(table: Table) -> None:
    """Center and bold every paragraph in every cell"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = CENTER
                for run in paragraph.runs:
                    run.bold = True


def _fill_titles(
    cell: _Cell, title: str | None, subtitle: str | None, bold: bool = True
):
    """Fill the title cells of an image(s) container table"""
    title_par: Paragraph = cell.paragraphs[0]
    title_run: Run = title_par.add_run(title)
    title_run.bold = bold
    if subtitle is not None:
        title_run.add_break()
        subtitle_run: Run = title_par.add_run(subtitle)
        subtitle_run.bold = False
        subtitle_run.italic = True


class Image:
    """A single image with an optional title"""

    def __init__(
        self,
        file: Path,
        title: str | None = None,
        subtitle: str | None = None,
        width: int = 4,
    ):
        self._title: str | None = title
        self._subtitle: str | None = subtitle
        self._file: Path = file
        self._width = width

    @classmethod
    def from_yaml(cls, loader: TemplateLoader, node: MappingNode) -> Self:
        image_cfg: dict[str, Any] = loader.construct_mapping(node, deep=True)
        return cls(
            file=loader.imports["image"] / image_cfg.pop("id"), **image_cfg
        )

    def render(self, document: Document) -> Document:
        # add container table
        container_table: Table = document.add_table(2, 1)
        container_table.alignment = CENTER
        _style_container_table(container_table)

        # set titles
        title_cell: _Cell = container_table.cell(0, 0)
        _fill_titles(title_cell, title=self._title, subtitle=self._subtitle)

        # fill container table
        content_cell: _Cell = container_table.cell(1, 0)
        content_par: Paragraph = content_cell.paragraphs[0]
        content_run: Run = content_par.add_run()
        content_run.add_picture(str(self._file), width=Inches(self._width))

        return document


class Images:
    """Multiple images side by side"""

    def __init__(
        self,
        id: str,
        images: list[Image],
        title: str | None = None,
        subtitle: str | None = None,
        width: int = 3,
    ) -> None:
        self._id: str = id
        self._images: list[Image] = images
        self._title: str = title
        self._subtitle: str = subtitle
        self._width: int = width

    @classmethod
    def from_yaml(cls, loader: TemplateLoader, node: MappingNode) -> Self:
        images_cfg: dict[str, Any] = loader.construct_mapping(node, deep=True)
        images_dict: dict[str, Image] = {
            id: image
            for (id, image) in images_cfg.items()
            if isinstance(image, Image)
        }
        other_dict: dict[str, Any] = {
            k: v for (k, v) in images_cfg.items() if k not in images_dict
        }
        return cls(
            id=other_dict.pop("id"),
            images=list(images_dict.values()),
            **other_dict
        )

    def render(self, document: Document) -> Document:
        # add container table
        columns: int = len(self._images)
        rows: int = 2 if columns == 1 else 3
        container_table: Table = document.add_table(rows, columns)
        container_table.alignment = CENTER
        _style_container_table(container_table)

        # merge title cells if necessary
        title_cell: _Cell = container_table.cell(0, 0)
        for idx, _ in enumerate(container_table.columns):
            title_cell.merge(container_table.cell(0, idx))

        # fill container table titles
        _fill_titles(title_cell, title=self._title, subtitle=self._subtitle)
        # fill container table subtitles
        for idx, subtitle_cell in enumerate(container_table.rows[1].cells):
            _fill_titles(
                subtitle_cell,
                title=self._images[idx]._title,
                subtitle=self._images[idx]._subtitle,
                bold=False,
            )

        # fill container table plots
        for idx, content_cell in enumerate(container_table.rows[2].cells):
            content_par: Paragraph = content_cell.paragraphs[0]
            content_run: Run = content_par.add_run()
            content_run.add_picture(
                str(self._images[idx]._file), width=Inches(self._width)
            )
        return document
