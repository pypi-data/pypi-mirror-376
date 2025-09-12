from pathlib import Path
from typing import Any, Self

import numpy as np
import pandas as pd
from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as TableDocx
from docx.table import _Cell as TableCell
from yaml import MappingNode

from tesorotools.utils.config import read_config
from tesorotools.utils.globals import PLOT_CONFIG_FILE
from tesorotools.utils.template import TemplateLoader

RENDER_CONFIG: dict[str, Any] = read_config(PLOT_CONFIG_FILE)["table"]

TEXTO_TABLAS = 9

CENTER = WD_ALIGN_PARAGRAPH.CENTER


def _set_cell_border(cell: TableCell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def _style_horizontal_blocks_header(cell: TableCell):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.size = Pt(12)


def _horizontal_blocks_header(columns: pd.MultiIndex, table_docx: TableDocx):
    column_counter: int = 1
    blocks: list[str] = list(columns.get_level_values(level=0).unique())
    for block in blocks:
        cell: TableCell = table_docx.cell(0, column_counter)
        columns_to_merge: int = len(
            columns[columns.get_level_values(level=0) == block]
        )
        for _ in range(columns_to_merge - 1):
            column_counter = column_counter + 1
            cell.merge(table_docx.cell(0, column_counter))
        column_counter = column_counter + 1
        cell.text = block
        _style_horizontal_blocks_header(cell)


def _style_column_names(cell: TableCell):
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fill_column_names(
    table: pd.DataFrame, table_docx: TableDocx, horizontal: bool
):
    if horizontal:
        start_row: int = 1
        _horizontal_blocks_header(table.columns, table_docx)
        columns: np.ndarray = table.columns.get_level_values(level=1).values
    else:
        start_row: int = 0
        columns: np.ndarray = table.columns.values

    for idx, column_name in enumerate(columns, start=1):
        cell: TableCell = table_docx.cell(start_row, idx)
        cell.text = column_name
        _style_column_names(cell)


def _style_index_names(cell: TableCell):
    cell.paragraphs[0].runs[0].font.size = Pt(TEXTO_TABLAS)
    cell.paragraphs[0].runs[0].font.bold = True
    cell.width = Inches(1)


def _fill_index_names(
    index: pd.Index | pd.MultiIndex, table_docx: TableDocx, horizontal: bool
):
    start_row: int = 2 if horizontal else 1

    index_names: pd.Index = (
        index
        if (horizontal or (not isinstance(index, pd.MultiIndex)))
        else index.get_level_values(level=1)
    )

    for idx, name in enumerate(index_names, start=start_row):
        cell: TableCell = table_docx.cell(idx, 0)
        print(name)
        cell.text = name
        _style_index_names(cell)


# we only separate blocks in vertically stacked tables
def _separate_blocks(index: pd.MultiIndex, table_docx: TableDocx):
    blocks: list[str] = list(index.get_level_values(level=0).unique())
    previous_rows = 0
    for block in blocks[:-1]:
        block_size: int = len(index[index.get_level_values(level=0) == block])
        for cell in table_docx.rows[block_size + previous_rows].cells:
            _separate_cell(cell)
        previous_rows += block_size


def _separate_cell(cell: TableCell):
    _set_cell_border(
        cell,
        bottom={
            "sz": 1,
            "val": "double",
            "color": "#000000",
            "space": 2,
        },
    )


def _is_bright(hex_color):
    red = int(hex_color[:2], 16)
    green = int(hex_color[2:4], 16)
    blue = int(hex_color[4:], 16)
    luminance = 0.2126 * red + 0.7152 * green + 0.0722 * blue
    return luminance > 180


def _shade_cell(cell: TableCell, hex_color: str):
    bright = _is_bright(hex_color)
    shading_element = parse_xml(
        r'<w:shd {} w:fill="{hex_color}"/>'.format(
            nsdecls("w"), hex_color=hex_color
        )
    )
    cell._tc.get_or_add_tcPr().append(shading_element)
    if bright:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
    else:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)


def _style_content(cell: TableCell):
    cell.paragraphs[0].runs[0].font.size = Pt(TEXTO_TABLAS)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fill_content(
    table: pd.DataFrame,
    color_table: pd.DataFrame,
    shade_table: pd.DataFrame,
    table_docx: TableDocx,
    horizontal: bool,
):
    start_row: int = 2 if horizontal else 1
    values = table.values
    for (x, y), value in np.ndenumerate(values):
        cell: TableCell = table_docx.cell(x + start_row, y + 1)
        cell.text = value if value is not None else ""
        color: str | None = (
            color_table.values[x, y] if color_table is not None else None
        )
        shade: str | None = (
            shade_table.values[x, y] if shade_table is not None else None
        )
        if color is not None:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(
                color
            )
        if shade is not None:
            _shade_cell(cell, shade)
        _style_content(cell)


def _style_table(table_docx: TableDocx):
    table_docx.style = RENDER_CONFIG.get("style", None)
    table_docx.autofit = RENDER_CONFIG["autofit"]


def render_table(
    table: pd.DataFrame,
    color_table: pd.DataFrame,
    shade_table: pd.DataFrame,
    document: Document,
    block_sep: bool,
    **kwargs,
) -> TableDocx:

    horizontal: bool = isinstance(table.columns, pd.MultiIndex)
    table_docx: TableDocx = document.add_table(
        rows=len(table.index) + table.columns.nlevels,
        cols=len(table.columns) + 1,
    )

    _style_table(table_docx)
    _fill_column_names(table, table_docx, horizontal)
    _fill_index_names(
        index=table.index, table_docx=table_docx, horizontal=horizontal
    )
    if block_sep:
        _separate_blocks(table.index, table_docx)
    _fill_content(table, color_table, shade_table, table_docx, horizontal)
    table_docx.alignment = WD_TABLE_ALIGNMENT.CENTER
    return document


class Table:
    """A rendered table in the document"""

    def __init__(
        self,
        data_file: Path | None = None,
        color_file: Path | None = None,
        shade_file: Path | None = None,
        block_sep: bool = False,
        title: str | None = None,
        columns: list[str] | None = None,
    ):
        if (
            (data_file is None)
            and (color_file is None)
            and (shade_file is None)
        ):
            raise ValueError("At least a piece of data should be given")
        self._data: pd.DataFrame | None = (
            pd.read_feather(data_file) if data_file is not None else None
        )
        self._color: pd.DataFrame | None = (
            pd.read_feather(color_file) if color_file is not None else None
        )
        self._shade: pd.DataFrame | None = (
            pd.read_feather(shade_file) if shade_file is not None else None
        )

        if columns is not None:
            self._data = self._data[columns]
            self._color = self._color[columns]
            self._shade = self._shade[columns]

        self._title: str | None = title
        self._block_sep: bool = block_sep

    @classmethod
    def from_yaml(cls, loader: TemplateLoader, node: MappingNode) -> Self:
        table_cfg: dict[str, Any] = loader.construct_mapping(node, deep=True)
        root_path: Path = loader.imports["table"]
        file_prefix: str = table_cfg.pop("id")

        # the data file has an optional "_data" suffix
        # if a file with the suffix exists, it will be preferred
        data_file_suffix: Path = root_path / f"{file_prefix}_data.feather"
        if data_file_suffix.exists():
            data_file: Path = data_file_suffix
        else:
            data_file: Path = root_path / f"{file_prefix}.feather"

        color_file: Path = root_path / f"{file_prefix}_color.feather"
        shade_file: Path = root_path / f"{file_prefix}_shade.feather"
        return cls(
            data_file,
            color_file=color_file if color_file.exists() else None,
            shade_file=shade_file if shade_file.exists() else None,
            **table_cfg,
        )

    def render(self, document: Document) -> Document:
        heading = document.add_heading(self._title, level=2)
        heading.alignment = CENTER
        heading.runs[0].font.size = Pt(10)

        print(self._data)
        print(self._data.index)

        render_table(
            self._data,
            self._color,
            self._shade,
            document,
            block_sep=self._block_sep,
        )
        return document
