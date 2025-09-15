"""
Document processing utilities for handling .docx files with template substitution.
"""

import os
import re
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

from .logger import get_logger
from .template_processor import TemplateProcessor


class DocumentProcessor:
    """Processes .docx documents with template substitution and formatting preservation."""
    
    def __init__(self):
        self.logger = get_logger(__name__)
        self.template_processor = TemplateProcessor()
    
    def process_docx_template(self, docx_path: str, recipient_data: Dict[str, str]) -> Document:
        """
        Process a .docx template by substituting {{variables}} with recipient data.
        
        Args:
            docx_path: Path to the .docx template file
            recipient_data: Dictionary containing recipient data for substitution
            
        Returns:
            Document object with processed content
        """
        try:
            # Load the document
            doc = Document(docx_path)
            
            # Process all paragraphs
            for paragraph in doc.paragraphs:
                self._process_paragraph(paragraph, recipient_data)
            
            # Process tables
            for table in doc.tables:
                self._process_table(table, recipient_data)
            
            # Process headers and footers
            for section in doc.sections:
                self._process_header_footer(section.header, recipient_data)
                self._process_header_footer(section.footer, recipient_data)
            
            return doc
            
        except Exception as e:
            self.logger.error(f"Error processing .docx template: {e}")
            raise
    
    def docx_to_html(self, doc: Document) -> str:
        """
        Convert a .docx document to HTML format for email.
        
        Args:
            doc: Document object to convert
            
        Returns:
            HTML string representation of the document
        """
        html_parts = ['<html><body>']
        
        try:
            for paragraph in doc.paragraphs:
                html_parts.append(self._paragraph_to_html(paragraph))
            
            for table in doc.tables:
                html_parts.append(self._table_to_html(table))
            
            html_parts.append('</body></html>')
            
            return '\n'.join(html_parts)
            
        except Exception as e:
            self.logger.error(f"Error converting .docx to HTML: {e}")
            # Fallback to plain text
            return self._docx_to_plain_text(doc)
    
    def docx_to_plain_text(self, doc: Document) -> str:
        """
        Convert a .docx document to plain text.
        
        Args:
            doc: Document object to convert
            
        Returns:
            Plain text representation of the document
        """
        return self._docx_to_plain_text(doc)
    
    def validate_docx_template(self, docx_path: str, available_fields: List[str]):
        """
        Validate .docx template and return validation result compatible with template processor.
        
        Args:
            docx_path: Path to the .docx template file
            available_fields: List of available field names from CSV
            
        Returns:
            Validation result object with is_valid and errors attributes
        """
        from collections import namedtuple
        ValidationResult = namedtuple('ValidationResult', ['is_valid', 'errors'])
        
        errors = []
        
        try:
            doc = Document(docx_path)
            
            # Extract all text content
            all_text = []
            
            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            
            # Get text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        all_text.append(paragraph.text)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        all_text.append(paragraph.text)
            
            # Find all template references
            full_text = '\n'.join(all_text)
            template_refs = self.template_processor._extract_template_references(full_text)
            
            # Check if all references are valid
            available_fields_set = set(available_fields)
            for ref in template_refs:
                if ref not in available_fields_set:
                    errors.append(f"Template reference '{{{{{ref}}}}}' in {os.path.basename(docx_path)} not found in CSV columns")
            
            is_valid = len(errors) == 0
            return ValidationResult(is_valid=is_valid, errors=errors)
            
        except Exception as e:
            errors.append(f"Error validating .docx template: {e}")
            return ValidationResult(is_valid=False, errors=errors)
    
    def _process_paragraph(self, paragraph, recipient_data: Dict[str, str]):
        """Process a single paragraph for template substitution while preserving run formatting."""
        if not paragraph.text:
            return
        
        # Check if paragraph contains any template variables
        if '{{' not in paragraph.text:
            return
        
        # Process each run individually to preserve formatting
        for run in paragraph.runs:
            if run.text and '{{' in run.text:
                # Process template variables in this run
                original_text = run.text
                processed_text = self.template_processor.process_template(original_text, recipient_data)
                
                if original_text != processed_text:
                    # Update the run text while preserving all formatting
                    run.text = processed_text
        
        # Handle cross-run variables (variables that span multiple runs)
        self._handle_cross_run_variables(paragraph, recipient_data)
    
    def _handle_cross_run_variables(self, paragraph, recipient_data: Dict[str, str]):
        """Handle template variables that span across multiple runs."""
        # Get the full paragraph text to check for cross-run variables
        full_text = paragraph.text
        
        # Find all template variables in the full text
        import re
        pattern = r'\{\{([^}]+)\}\}'
        variables_in_full_text = set(re.findall(pattern, full_text))
        
        if not variables_in_full_text:
            return
        
        # Check each variable to see if it's split across runs
        for var_name in variables_in_full_text:
            var_pattern = '{{' + var_name + '}}'
            
            # Check if this variable is split across runs
            if self._is_variable_split_across_runs(paragraph, var_pattern):
                self._merge_split_variable(paragraph, var_pattern, recipient_data)
    
    def _is_variable_split_across_runs(self, paragraph, variable_pattern: str) -> bool:
        """Check if a template variable is split across multiple runs."""
        # If any single run contains the complete variable, it's not split
        for run in paragraph.runs:
            if variable_pattern in run.text:
                return False
        
        # Check if the variable exists in the full paragraph text but not in any single run
        return variable_pattern in paragraph.text
    
    def _merge_split_variable(self, paragraph, variable_pattern: str, recipient_data: Dict[str, str]):
        """Merge a template variable that's split across multiple runs."""
        # This is a complex case where a variable like {{name}} might be split as:
        # Run 1: "Hello {{"
        # Run 2: "name"  
        # Run 3: "}} from"
        
        # Find the runs that contain parts of the variable
        var_start = variable_pattern[:2]  # "{{"
        var_end = variable_pattern[-2:]   # "}}"
        var_content = variable_pattern[2:-2]  # "name"
        
        # Find start and end run indices
        start_run_idx = None
        end_run_idx = None
        
        for i, run in enumerate(paragraph.runs):
            if var_start in run.text and start_run_idx is None:
                start_run_idx = i
            if var_end in run.text:
                end_run_idx = i
        
        if start_run_idx is not None and end_run_idx is not None:
            # Process the variable
            replacement_value = recipient_data.get(var_content.strip(), f"{{{{MISSING: {var_content.strip()}}}}}")
            
            # Reconstruct the text across the affected runs
            self._reconstruct_cross_run_text(paragraph, start_run_idx, end_run_idx, 
                                           variable_pattern, replacement_value)
    
    def _reconstruct_cross_run_text(self, paragraph, start_idx: int, end_idx: int, 
                                  variable_pattern: str, replacement_value: str):
        """Reconstruct text across multiple runs while preserving formatting."""
        if start_idx == end_idx:
            # Variable is within a single run
            run = paragraph.runs[start_idx]
            run.text = run.text.replace(variable_pattern, replacement_value)
            return
        
        # Variable spans multiple runs - this is the complex case
        # We'll merge the content but preserve the formatting of the first run
        
        # Collect text from all affected runs
        combined_text = ""
        for i in range(start_idx, end_idx + 1):
            combined_text += paragraph.runs[i].text
        
        # Replace the variable in the combined text
        new_combined_text = combined_text.replace(variable_pattern, replacement_value)
        
        # Put the new text in the first run and clear the others
        paragraph.runs[start_idx].text = new_combined_text
        
        # Clear the other runs (but preserve them for formatting structure)
        for i in range(start_idx + 1, end_idx + 1):
            paragraph.runs[i].text = ""
    
    def _process_table(self, table, recipient_data: Dict[str, str]):
        """Process all cells in a table for template substitution."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, recipient_data)
    
    def _process_header_footer(self, header_footer, recipient_data: Dict[str, str]):
        """Process header or footer for template substitution."""
        if header_footer:
            for paragraph in header_footer.paragraphs:
                self._process_paragraph(paragraph, recipient_data)
    
    def _paragraph_to_html(self, paragraph) -> str:
        """Convert a paragraph to HTML."""
        if not paragraph.text.strip():
            return '<br>'
        
        # Basic paragraph conversion
        html = '<p'
        
        # Handle alignment
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            html += ' style="text-align: center;"'
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            html += ' style="text-align: right;"'
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            html += ' style="text-align: justify;"'
        
        html += '>'
        
        # Process runs for formatting
        for run in paragraph.runs:
            run_html = run.text
            
            if run.bold:
                run_html = f'<strong>{run_html}</strong>'
            if run.italic:
                run_html = f'<em>{run_html}</em>'
            if run.underline:
                run_html = f'<u>{run_html}</u>'
            
            # Handle font size
            if run.font.size:
                size_pt = run.font.size.pt
                run_html = f'<span style="font-size: {size_pt}pt;">{run_html}</span>'
            
            # Handle font color
            if run.font.color.rgb:
                color = run.font.color.rgb
                run_html = f'<span style="color: #{color};">{run_html}</span>'
            
            html += run_html
        
        html += '</p>'
        return html
    
    def _table_to_html(self, table) -> str:
        """Convert a table to HTML."""
        html = ['<table border="1" style="border-collapse: collapse;">']
        
        for row in table.rows:
            html.append('<tr>')
            for cell in row.cells:
                cell_html = '<td>'
                for paragraph in cell.paragraphs:
                    cell_html += self._paragraph_to_html(paragraph)
                cell_html += '</td>'
                html.append(cell_html)
            html.append('</tr>')
        
        html.append('</table>')
        return '\n'.join(html)
    
    def _docx_to_plain_text(self, doc: Document) -> str:
        """Convert document to plain text as fallback."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text for p in cell.paragraphs if p.text.strip())
                    row_text.append(cell_text)
                if any(row_text):
                    text_parts.append(' | '.join(row_text))
        
        return '\n\n'.join(text_parts)
    
    def render_docx_to_image(self, doc: Document, width: int = 400, height: int = 300) -> bytes:
        """
        Render .docx content to an image for picture fusion with formatting preservation.
        
        Args:
            doc: Document to render
            width: Target width in pixels
            height: Target height in pixels
            
        Returns:
            Image bytes in PNG format
        """
        from PIL import Image, ImageDraw, ImageFont
        from io import BytesIO
        
        # Create a transparent background image
        img = Image.new('RGBA', (width, height), color=(255, 255, 255, 0))  # Transparent white
        draw = ImageDraw.Draw(img)
        
        # Font cache for different sizes and styles
        font_cache = {}
        
        def get_font(font_name: str = 'Arial', size: int = 12, bold: bool = False, italic: bool = False):
            """Get font with enhanced font loading system."""
            key = (font_name, size, bold, italic)
            if key in font_cache:
                return font_cache[key]
            
            try:
                # Create instruction dict for the enhanced font system
                shape_parts = []
                if bold:
                    shape_parts.append('bold')
                if italic:
                    shape_parts.append('italic')
                shape = ' '.join(shape_parts)
                
                instruction = {
                    'font': font_name,
                    'shape': shape,
                    'size': str(size)
                }
                
                # Use the enhanced font loading system from picture generator
                font = _get_enhanced_font(instruction, size)
                font_cache[key] = font
                return font
                
            except Exception as e:
                # Fallback to default font
                font = ImageFont.load_default()
                font_cache[key] = font
                return font
        
        def _get_enhanced_font(instruction: dict, size: int):
            """Enhanced font loading with cross-platform support."""
            # Get font family from font parameter
            font_family = instruction.get('font', 'arial').strip()
            
            # Get font style from shape
            font_style = instruction.get('shape', '').strip().lower()
            
            # Parse font style
            is_bold = 'bold' in font_style
            is_italic = 'italic' in font_style
            
            # Normalize font family name
            font_family_lower = font_family.lower()
            
            # Cross-platform font loading
            font = _load_cross_platform_font(font_family_lower, size, is_bold, is_italic)
            if font:
                return font
            
            # Fallback to default font
            return ImageFont.load_default()
        
        def _load_cross_platform_font(font_family: str, size: int, is_bold: bool, is_italic: bool):
            """Load fonts across Mac, Windows, and Linux platforms."""
            import platform
            system = platform.system()
            
            # Try platform-specific font loading
            if system == "Darwin":  # macOS
                return _load_macos_font(font_family, size, is_bold, is_italic)
            elif system == "Windows":
                return _load_windows_font(font_family, size, is_bold, is_italic)
            else:  # Linux and others
                return _load_linux_font(font_family, size, is_bold, is_italic)
        
        def _load_macos_font(font_family: str, size: int, is_bold: bool, is_italic: bool):
            """Load fonts on macOS."""
            # macOS system fonts with .ttc support
            macos_fonts = {
                'arial': 'ArialHB.ttc',
                'helvetica': 'Helvetica.ttc',
                'courier': 'Courier.ttc',
                'courier new': 'Courier.ttc',
                'times': 'Times.ttc',
                'times new roman': 'Times.ttc',
                'georgia': 'Georgia.ttc',
                'verdana': 'Verdana.ttc',
                'monaco': 'Monaco.ttf',
                'menlo': 'Menlo.ttc',
            }
            
            font_file = macos_fonts.get(font_family)
            if not font_file:
                return None
            
            font_path = f"/System/Library/Fonts/{font_file}"
            if not os.path.exists(font_path):
                return None
            
            try:
                if font_path.endswith('.ttc'):
                    # For .ttc files, try different indices for styles
                    if is_bold and is_italic:
                        indices = [3, 1, 0]  # Bold Italic, Bold, Regular
                    elif is_bold:
                        indices = [1, 0]  # Bold, Regular
                    elif is_italic:
                        indices = [2, 0]  # Italic, Regular
                    else:
                        indices = [0]  # Regular
                    
                    for index in indices:
                        try:
                            return ImageFont.truetype(font_path, size, index=index)
                        except Exception:
                            continue
                else:
                    return ImageFont.truetype(font_path, size)
            except Exception:
                return None
        
        def _load_windows_font(font_family: str, size: int, is_bold: bool, is_italic: bool):
            """Load fonts on Windows."""
            # Windows font files
            windows_fonts = {
                'arial': {
                    'regular': 'arial.ttf',
                    'bold': 'arialbd.ttf',
                    'italic': 'ariali.ttf',
                    'bold_italic': 'arialbi.ttf'
                },
                'helvetica': {
                    'regular': 'arial.ttf',  # Fallback to Arial on Windows
                    'bold': 'arialbd.ttf',
                    'italic': 'ariali.ttf',
                    'bold_italic': 'arialbi.ttf'
                },
                'courier': {
                    'regular': 'cour.ttf',
                    'bold': 'courbd.ttf',
                    'italic': 'couri.ttf',
                    'bold_italic': 'courbi.ttf'
                },
                'courier new': {
                    'regular': 'cour.ttf',
                    'bold': 'courbd.ttf',
                    'italic': 'couri.ttf',
                    'bold_italic': 'courbi.ttf'
                },
                'times': {
                    'regular': 'times.ttf',
                    'bold': 'timesbd.ttf',
                    'italic': 'timesi.ttf',
                    'bold_italic': 'timesbi.ttf'
                },
                'times new roman': {
                    'regular': 'times.ttf',
                    'bold': 'timesbd.ttf',
                    'italic': 'timesi.ttf',
                    'bold_italic': 'timesbi.ttf'
                },
                'georgia': {
                    'regular': 'georgia.ttf',
                    'bold': 'georgiab.ttf',
                    'italic': 'georgiai.ttf',
                    'bold_italic': 'georgiaz.ttf'
                },
                'verdana': {
                    'regular': 'verdana.ttf',
                    'bold': 'verdanab.ttf',
                    'italic': 'verdanai.ttf',
                    'bold_italic': 'verdanaz.ttf'
                },
                'monaco': {
                    'regular': 'consola.ttf',  # Fallback to Consolas
                    'bold': 'consolab.ttf',
                    'italic': 'consolai.ttf',
                    'bold_italic': 'consolaz.ttf'
                }
            }
            
            font_variants = windows_fonts.get(font_family)
            if not font_variants:
                return None
            
            # Select appropriate variant
            if is_bold and is_italic:
                font_file = font_variants.get('bold_italic', font_variants['regular'])
            elif is_bold:
                font_file = font_variants.get('bold', font_variants['regular'])
            elif is_italic:
                font_file = font_variants.get('italic', font_variants['regular'])
            else:
                font_file = font_variants['regular']
            
            # Try Windows font directories
            font_dirs = [
                "C:/Windows/Fonts/",
                "C:/WINDOWS/Fonts/",
                os.path.expanduser("~/AppData/Local/Microsoft/Windows/Fonts/")
            ]
            
            for font_dir in font_dirs:
                font_path = os.path.join(font_dir, font_file)
                if os.path.exists(font_path):
                    try:
                        return ImageFont.truetype(font_path, size)
                    except Exception:
                        continue
            
            return None
        
        def _load_linux_font(font_family: str, size: int, is_bold: bool, is_italic: bool):
            """Load fonts on Linux."""
            # Linux font mappings (using common distributions)
            linux_fonts = {
                'arial': {
                    'regular': ['DejaVuSans.ttf', 'LiberationSans-Regular.ttf'],
                    'bold': ['DejaVuSans-Bold.ttf', 'LiberationSans-Bold.ttf'],
                    'italic': ['DejaVuSans-Oblique.ttf', 'LiberationSans-Italic.ttf'],
                    'bold_italic': ['DejaVuSans-BoldOblique.ttf', 'LiberationSans-BoldItalic.ttf']
                },
                'helvetica': {
                    'regular': ['DejaVuSans.ttf', 'LiberationSans-Regular.ttf'],
                    'bold': ['DejaVuSans-Bold.ttf', 'LiberationSans-Bold.ttf'],
                    'italic': ['DejaVuSans-Oblique.ttf', 'LiberationSans-Italic.ttf'],
                    'bold_italic': ['DejaVuSans-BoldOblique.ttf', 'LiberationSans-BoldItalic.ttf']
                },
                'courier': {
                    'regular': ['DejaVuSansMono.ttf', 'LiberationMono-Regular.ttf'],
                    'bold': ['DejaVuSansMono-Bold.ttf', 'LiberationMono-Bold.ttf'],
                    'italic': ['DejaVuSansMono-Oblique.ttf', 'LiberationMono-Italic.ttf'],
                    'bold_italic': ['DejaVuSansMono-BoldOblique.ttf', 'LiberationMono-BoldItalic.ttf']
                },
                'courier new': {
                    'regular': ['DejaVuSansMono.ttf', 'LiberationMono-Regular.ttf'],
                    'bold': ['DejaVuSansMono-Bold.ttf', 'LiberationMono-Bold.ttf'],
                    'italic': ['DejaVuSansMono-Oblique.ttf', 'LiberationMono-Italic.ttf'],
                    'bold_italic': ['DejaVuSansMono-BoldOblique.ttf', 'LiberationMono-BoldItalic.ttf']
                },
                'times': {
                    'regular': ['DejaVuSerif.ttf', 'LiberationSerif-Regular.ttf'],
                    'bold': ['DejaVuSerif-Bold.ttf', 'LiberationSerif-Bold.ttf'],
                    'italic': ['DejaVuSerif-Italic.ttf', 'LiberationSerif-Italic.ttf'],
                    'bold_italic': ['DejaVuSerif-BoldItalic.ttf', 'LiberationSerif-BoldItalic.ttf']
                },
                'times new roman': {
                    'regular': ['DejaVuSerif.ttf', 'LiberationSerif-Regular.ttf'],
                    'bold': ['DejaVuSerif-Bold.ttf', 'LiberationSerif-Bold.ttf'],
                    'italic': ['DejaVuSerif-Italic.ttf', 'LiberationSerif-Italic.ttf'],
                    'bold_italic': ['DejaVuSerif-BoldItalic.ttf', 'LiberationSerif-BoldItalic.ttf']
                },
                'monaco': {
                    'regular': ['DejaVuSansMono.ttf', 'LiberationMono-Regular.ttf'],
                    'bold': ['DejaVuSansMono-Bold.ttf', 'LiberationMono-Bold.ttf'],
                    'italic': ['DejaVuSansMono-Oblique.ttf', 'LiberationMono-Italic.ttf'],
                    'bold_italic': ['DejaVuSansMono-BoldOblique.ttf', 'LiberationMono-BoldItalic.ttf']
                }
            }
            
            font_variants = linux_fonts.get(font_family)
            if not font_variants:
                return None
            
            # Select appropriate variant
            if is_bold and is_italic:
                font_files = font_variants.get('bold_italic', font_variants['regular'])
            elif is_bold:
                font_files = font_variants.get('bold', font_variants['regular'])
            elif is_italic:
                font_files = font_variants.get('italic', font_variants['regular'])
            else:
                font_files = font_variants['regular']
            
            # Try Linux font directories
            font_dirs = [
                "/usr/share/fonts/truetype/dejavu/",
                "/usr/share/fonts/truetype/liberation/",
                "/usr/share/fonts/TTF/",
                "/usr/share/fonts/truetype/",
                "/usr/local/share/fonts/",
                os.path.expanduser("~/.fonts/"),
                os.path.expanduser("~/.local/share/fonts/")
            ]
            
            for font_file in font_files:
                for font_dir in font_dirs:
                    font_path = os.path.join(font_dir, font_file)
                    if os.path.exists(font_path):
                        try:
                            return ImageFont.truetype(font_path, size)
                        except Exception:
                            continue
            
            return None
        
        def parse_color(color_obj):
            """Parse docx color to PIL color."""
            if color_obj and color_obj.rgb:
                # Convert RGB to hex string
                rgb_int = int(str(color_obj.rgb), 16)
                r = (rgb_int >> 16) & 255
                g = (rgb_int >> 8) & 255
                b = rgb_int & 255
                return (r, g, b)
            return (0, 0, 0)  # Default to black
        
        def wrap_text(text: str, font, max_width: int):
            """Wrap text to fit within max_width."""
            words = text.split()
            lines = []
            current_line = []
            
            for word in words:
                test_line = ' '.join(current_line + [word])
                bbox = draw.textbbox((0, 0), test_line, font=font)
                text_width = bbox[2] - bbox[0]
                
                if text_width <= max_width:
                    current_line.append(word)
                else:
                    if current_line:
                        lines.append(' '.join(current_line))
                        current_line = [word]
                    else:
                        # Single word is too long, add it anyway
                        lines.append(word)
            
            if current_line:
                lines.append(' '.join(current_line))
            
            return lines
        
        # Rendering parameters
        y_offset = 10
        margin_left = 10
        margin_right = 10
        max_text_width = width - margin_left - margin_right
        
        # Process each paragraph with formatting
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                # Empty paragraph - add some spacing
                y_offset += 10
                continue
            
            # Handle paragraph alignment
            para_alignment = paragraph.alignment
            
            # Process runs within the paragraph to preserve formatting
            paragraph_lines = []
            current_line_runs = []
            
            for run in paragraph.runs:
                if not run.text:
                    continue
                
                # Get run formatting
                font_size = 12  # Default size
                if run.font.size:
                    font_size = int(run.font.size.pt)
                
                is_bold = run.bold if run.bold is not None else False
                is_italic = run.italic if run.italic is not None else False
                color = parse_color(run.font.color)
                
                # Get font family name
                font_name = 'Arial'  # Default
                if run.font.name:
                    font_name = run.font.name
                
                # Get appropriate font using enhanced system
                font = get_font(font_name, font_size, is_bold, is_italic)
                
                # Wrap the run text
                run_lines = wrap_text(run.text, font, max_text_width)
                
                for i, line in enumerate(run_lines):
                    if i == 0 and current_line_runs:
                        # Continue current line
                        current_line_runs.append((line, font, color))
                    else:
                        # Start new line
                        if current_line_runs:
                            paragraph_lines.append(current_line_runs)
                        current_line_runs = [(line, font, color)]
            
            # Add the last line
            if current_line_runs:
                paragraph_lines.append(current_line_runs)
            
            # Render the paragraph lines
            for line_runs in paragraph_lines:
                if y_offset >= height - 20:  # Stop if we're near the bottom
                    break
                
                # Calculate line height based on largest font in the line
                line_height = max(20, max(font.size for _, font, _ in line_runs) + 4)
                
                # Calculate total line width for alignment
                total_line_width = 0
                for text, font, color in line_runs:
                    bbox = draw.textbbox((0, 0), text, font=font)
                    total_line_width += bbox[2] - bbox[0]
                
                # Calculate starting x position based on alignment
                if para_alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    x_start = margin_left + (max_text_width - total_line_width) // 2
                elif para_alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    x_start = width - margin_right - total_line_width
                else:  # LEFT or JUSTIFY
                    x_start = margin_left
                
                # Ensure x_start is not negative
                x_start = max(margin_left, x_start)
                
                # Render each run in the line
                x_current = x_start
                for text, font, color in line_runs:
                    if x_current < width - margin_right:
                        draw.text((x_current, y_offset), text, fill=color, font=font)
                        bbox = draw.textbbox((0, 0), text, font=font)
                        x_current += bbox[2] - bbox[0]
                
                y_offset += line_height
            
            # Add paragraph spacing
            y_offset += 5
        
        # Convert to bytes
        img_bytes = BytesIO()
        img.save(img_bytes, format='PNG')
        return img_bytes.getvalue()
    
    def convert_to_html(self, doc: Document) -> str:
        """
        Convert a processed Document to HTML format.
        
        Args:
            doc: Document object to convert
            
        Returns:
            HTML string representation of the document
        """
        html_parts = []
        
        # Convert paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                html_parts.append(self._paragraph_to_html(paragraph))
        
        # Convert tables
        for table in doc.tables:
            html_parts.append(self._table_to_html(table))
        
        return '\n'.join(html_parts)
    
    def convert_to_plain_text(self, doc: Document) -> str:
        """
        Convert a processed Document to plain text format.
        
        Args:
            doc: Document object to convert
            
        Returns:
            Plain text string representation of the document
        """
        return self._docx_to_plain_text(doc)
    
    def convert_to_html(self, doc: Document) -> str:
        """
        Convert a processed Document to HTML format.
        
        Args:
            doc: Document object to convert
            
        Returns:
            HTML string representation of the document
        """
        html_parts = []
        
        # Convert paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                html_parts.append(self._paragraph_to_html(paragraph))
        
        # Convert tables
        for table in doc.tables:
            html_parts.append(self._table_to_html(table))
        
        return '\n'.join(html_parts)
    
    def convert_to_plain_text(self, doc: Document) -> str:
        """
        Convert a processed Document to plain text format.
        
        Args:
            doc: Document object to convert
            
        Returns:
            Plain text string representation of the document
        """
        return self._docx_to_plain_text(doc)
