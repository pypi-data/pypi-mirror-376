"""
Word文档文本编辑工具函数

提供四个核心的文档文本编辑功能
"""

import os
from typing import Optional
from .utils import (
    ensure_docx_extension,
    check_file_writeable,
    find_and_replace_text,
    create_style
)


async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                     bold: Optional[bool] = None, italic: Optional[bool] = None, 
                     underline: Optional[bool] = None, color: Optional[str] = None,
                     font_size: Optional[int] = None, font_name: Optional[str] = None) -> str:
    """Format a specific range of text within a paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        start_pos: Start position within the paragraph text
        end_pos: End position within the paragraph text
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        paragraph_index = int(paragraph_index)
        start_pos = int(start_pos)
        end_pos = int(end_pos)
        if font_size is not None:
            font_size = int(font_size)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: paragraph_index, start_pos, end_pos, and font_size must be integers")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise ValueError(f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1}).")

        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text

        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
            raise ValueError(f"Invalid text positions. Paragraph has {len(text)} characters.")
        
        # Get the text to format
        target_text = text[start_pos:end_pos]
        
        # Clear existing runs and create three runs: before, target, after
        for run in paragraph.runs:
            run.clear()
        
        # Add text before target
        if start_pos > 0:
            run_before = paragraph.add_run(text[:start_pos])
        
        # Add target text with formatting
        run_target = paragraph.add_run(target_text)
        if bold is not None:
            run_target.bold = bold
        if italic is not None:
            run_target.italic = italic
        if underline is not None:
            run_target.underline = underline
        if color:
            # Color mapping for common colors
            color_map = {
                'red': RGBColor(255, 0, 0),
                'blue': RGBColor(0, 0, 255),
                'green': RGBColor(0, 128, 0),
                'black': RGBColor(0, 0, 0),
                'white': RGBColor(255, 255, 255),
                'yellow': RGBColor(255, 255, 0),
                'orange': RGBColor(255, 165, 0),
                'purple': RGBColor(128, 0, 128),
                'gray': RGBColor(128, 128, 128),
                'grey': RGBColor(128, 128, 128)
            }
            try:
                if color.lower() in color_map:
                    run_target.font.color.rgb = color_map[color.lower()]
                else:
                    run_target.font.color.rgb = RGBColor.from_string(color)
            except Exception:
                run_target.font.color.rgb = RGBColor(0, 0, 0)
        if font_size:
            run_target.font.size = Pt(font_size)
        if font_name:
            run_target.font.name = font_name
        
        # Add text after target
        if end_pos < len(text):
            run_after = paragraph.add_run(text[end_pos:])
        
        doc.save(filename)
        return f"Text '{target_text}' formatted successfully in paragraph {paragraph_index}."
    except Exception as e:
        raise RuntimeError(f"Failed to format text: {str(e)}")


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        raise RuntimeError(f"Failed to search and replace: {str(e)}")


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    # Ensure numeric parameters are the correct type
    try:
        paragraph_index = int(paragraph_index)
    except (ValueError, TypeError):
        raise ValueError("Invalid parameter: paragraph_index must be an integer")

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            raise ValueError(f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1}).")
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        raise RuntimeError(f"Failed to delete paragraph: {str(e)}")


async def create_custom_style(filename: str, style_name: str,
                             bold: Optional[bool] = None, italic: Optional[bool] = None,
                             font_size: Optional[int] = None, font_name: Optional[str] = None,
                             color: Optional[str] = None, base_style: Optional[str] = None) -> str:
    """Create a custom style in the document.

    Args:
        filename: Path to the Word document
        style_name: Name for the new style
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        font_size: Font size in points
        font_name: Font name/family
        color: Text color (e.g., 'red', 'blue')
        base_style: Optional existing style to base this on
    """
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        raise FileNotFoundError(f"Document {filename} does not exist")

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        raise PermissionError(f"Cannot modify document: {error_message}. Consider creating a copy first.")

    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        
        # Build font properties dictionary
        font_properties = {}
        if bold is not None:
            font_properties['bold'] = bold
        if italic is not None:
            font_properties['italic'] = italic
        if font_size is not None:
            font_properties['size'] = font_size
        if font_name is not None:
            font_properties['name'] = font_name
        if color is not None:
            font_properties['color'] = color

        # Create the style
        new_style = create_style(
            doc,
            style_name,
            WD_STYLE_TYPE.PARAGRAPH,
            base_style=base_style,
            font_properties=font_properties
        )

        doc.save(filename)
        return f"Style '{style_name}' created successfully."
    except Exception as e:
        raise RuntimeError(f"Failed to create style: {str(e)}")
