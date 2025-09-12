"""
Word文档文本编辑工具的辅助函数

提供文件处理、文本操作和样式创建相关的辅助功能
"""

import os
from typing import Optional, Dict, Any


def ensure_docx_extension(filename: str) -> str:
    """确保文件名有.docx扩展名
    
    Args:
        filename: 原始文件名
        
    Returns:
        带有.docx扩展名的文件名
    """
    if not filename.lower().endswith('.docx'):
        return filename + '.docx'
    return filename


def check_file_writeable(filename: str) -> tuple[bool, str]:
    """检查文件是否可写
    
    Args:
        filename: 文件路径
        
    Returns:
        (是否可写, 错误信息)
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(filename):
            return False, "File does not exist"
        
        # 检查文件权限
        if not os.access(filename, os.W_OK):
            return False, "File is not writable (permission denied)"
        
        # 检查文件是否被其他程序锁定
        try:
            # 尝试以写入模式打开文件
            with open(filename, 'r+b') as f:
                pass
        except PermissionError:
            return False, "File is currently open in another application"
        except IOError as e:
            return False, f"File access error: {str(e)}"
        
        return True, ""
    
    except Exception as e:
        return False, f"Error checking file: {str(e)}"


def find_and_replace_text(doc, old_text: str, new_text: str) -> int:
    """
    Find and replace text throughout the document, skipping Table of Contents (TOC) paragraphs.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        # Skip TOC paragraphs
        if para.style and para.style.name.startswith("TOC"):
            continue
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Skip TOC paragraphs in tables too
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count


def create_style(doc, style_name: str, style_type, base_style: Optional[str] = None, 
                font_properties: Optional[Dict[str, Any]] = None):
    """Create a custom style in the document.
    
    Args:
        doc: Document object
        style_name: Name for the new style
        style_type: Type of style (WD_STYLE_TYPE.PARAGRAPH, etc.)
        base_style: Optional existing style to base this on
        font_properties: Dictionary of font properties to apply
        
    Returns:
        The created style object
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    
    # Check if style already exists
    try:
        existing_style = doc.styles[style_name]
        return existing_style
    except KeyError:
        pass  # Style doesn't exist, create it
    
    # Create new style
    if base_style:
        try:
            base = doc.styles[base_style]
            new_style = doc.styles.add_style(style_name, style_type, base)
        except KeyError:
            # Base style doesn't exist, create without base
            new_style = doc.styles.add_style(style_name, style_type)
    else:
        new_style = doc.styles.add_style(style_name, style_type)
    
    # Apply font properties if provided
    if font_properties:
        font = new_style.font
        
        if 'bold' in font_properties:
            font.bold = font_properties['bold']
        
        if 'italic' in font_properties:
            font.italic = font_properties['italic']
        
        if 'size' in font_properties:
            font.size = Pt(font_properties['size'])
        
        if 'name' in font_properties:
            font.name = font_properties['name']
        
        if 'color' in font_properties:
            color = font_properties['color']
            # Color mapping for common colors
            color_map = {
                'red': RGBColor(255, 0, 0),
                'blue': RGBColor(0, 0, 255),
                'green': RGBColor(0, 128, 0),
                'black': RGBColor(0, 0, 0),
                'white': RGBColor(255, 255, 255),
                'yellow': RGBColor(255, 255, 0),
                'orange': RGBColor(255, 165, 0),
                'purple': RGBColor(128, 0, 128),
                'gray': RGBColor(128, 128, 128),
                'grey': RGBColor(128, 128, 128)
            }
            
            try:
                if color.lower() in color_map:
                    font.color.rgb = color_map[color.lower()]
                else:
                    # Try to parse as hex color
                    if color.startswith('#'):
                        color = color[1:]
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        font.color.rgb = RGBColor(r, g, b)
                    else:
                        # Default to black if color parsing fails
                        font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                # Default to black if color parsing fails
                font.color.rgb = RGBColor(0, 0, 0)
    
    return new_style


def get_paragraph_info(doc, paragraph_index: int) -> Dict[str, Any]:
    """获取段落的详细信息
    
    Args:
        doc: Document object
        paragraph_index: 段落索引
        
    Returns:
        包含段落信息的字典
    """
    try:
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"error": f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs."}
        
        paragraph = doc.paragraphs[paragraph_index]
        
        return {
            "index": paragraph_index,
            "text": paragraph.text,
            "text_length": len(paragraph.text),
            "style": paragraph.style.name if paragraph.style else "None",
            "runs_count": len(paragraph.runs),
            "runs": [{"text": run.text, "bold": run.bold, "italic": run.italic} for run in paragraph.runs]
        }
    
    except Exception as e:
        return {"error": f"Error getting paragraph info: {str(e)}"}


def validate_text_range(text: str, start_pos: int, end_pos: int) -> tuple[bool, str]:
    """验证文本范围是否有效
    
    Args:
        text: 文本内容
        start_pos: 开始位置
        end_pos: 结束位置
        
    Returns:
        (是否有效, 错误信息)
    """
    if start_pos < 0:
        return False, "Start position cannot be negative"
    
    if end_pos > len(text):
        return False, f"End position ({end_pos}) exceeds text length ({len(text)})"
    
    if start_pos >= end_pos:
        return False, "Start position must be less than end position"
    
    return True, ""
