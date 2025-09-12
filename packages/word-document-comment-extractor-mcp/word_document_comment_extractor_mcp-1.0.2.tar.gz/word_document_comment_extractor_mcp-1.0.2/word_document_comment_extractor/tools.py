"""
Word文档评论提取工具函数

提供三个核心的文档评论提取功能
"""

import os
import json
from typing import Dict, List, Optional, Any
from .utils import (
    ensure_docx_extension,
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph_by_index
)


async def get_all_comments(filename: str) -> str:
    """
    Extract all comments from a Word document.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string containing all comments with metadata
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')
    
    try:
        # Load the document with clear error on invalid .docx
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        comments = extract_all_comments(doc)
        return json.dumps({
            'success': True,
            'comments': comments,
            'total_comments': len(comments)
        }, indent=2)
    except Exception as e:
        raise RuntimeError(f"Failed to extract comments: {str(e)}")


async def get_comments_by_author(filename: str, author: str) -> str:
    """
    Extract comments from a specific author in a Word document.
    
    Args:
        filename: Path to the Word document
        author: Name of the comment author to filter by
        
    Returns:
        JSON string containing filtered comments
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')

    if not author or not author.strip():
        raise ValueError('Author name cannot be empty')
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise
        all_comments = extract_all_comments(doc)
        author_comments = filter_comments_by_author(all_comments, author)
        return json.dumps({
            'success': True,
            'author': author,
            'comments': author_comments,
            'total_comments': len(author_comments)
        }, indent=2)
    except Exception as e:
        raise RuntimeError(f"Failed to extract comments: {str(e)}")


async def get_comments_for_paragraph(filename: str, paragraph_index: int) -> str:
    """
    Extract comments for a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        
    Returns:
        JSON string containing comments for the specified paragraph
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')

    if paragraph_index < 0:
        raise ValueError('Paragraph index must be non-negative')
    
    try:
        try:
            doc = Document(filename)
        except Exception as open_error:
            try:
                from docx.opc.exceptions import PackageNotFoundError
            except Exception:
                PackageNotFoundError = tuple()  # type: ignore
            if isinstance(open_error, PackageNotFoundError):
                raise RuntimeError(
                    f"The file is not a valid .docx package or is corrupted: {os.path.abspath(filename)}. "
                    "Please ensure the document is a real .docx (not .doc) and not open/locked in another app."
                )
            raise

        if paragraph_index >= len(doc.paragraphs):
            raise ValueError(
                f"Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs."
            )

        all_comments = extract_all_comments(doc)
        para_comments = get_comments_for_paragraph_by_index(all_comments, paragraph_index)
        paragraph_text = doc.paragraphs[paragraph_index].text
        return json.dumps({
            'success': True,
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph_text,
            'comments': para_comments,
            'total_comments': len(para_comments)
        }, indent=2)
    except Exception as e:
        raise RuntimeError(f"Failed to extract comments: {str(e)}")
